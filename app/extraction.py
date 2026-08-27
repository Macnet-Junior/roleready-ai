import re
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile
from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader

MAX_FILE_BYTES = 5 * 1024 * 1024
MAX_TEXT_CHARS = 50_000
MAX_PDF_PAGES = 50
import re
import zipfile
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile
from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader

MAX_FILE_BYTES = 10 * 1024 * 1024
MAX_TEXT_CHARS = 50_000
MAX_PDF_PAGES = 50
ALLOWED = {
    ".pdf": {"application/pdf", "application/octet-stream"},
    ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/octet-stream"},
    ".md": {"text/markdown", "text/plain", "application/octet-stream"},
    ".markdown": {"text/markdown", "text/plain", "application/octet-stream"},
    ".txt": {"text/plain", "application/octet-stream"},
    ".zip": {"application/zip", "application/x-zip-compressed", "application/octet-stream"},
}

class ResumeExtractionError(ValueError): pass

def _optimize_extracted_text(text: str) -> str:
    """Token-aware text normalization: strips boilerplate, page footers, and excessive blank lines."""
    lines = [line.strip() for line in text.splitlines()]
    cleaned_lines = []
    page_footer_pattern = re.compile(r"^(page\s+\d+(\s+of\s+\d+)?|\d+\s*/\s*\d+)$", re.IGNORECASE)

    for line in lines:
        if page_footer_pattern.match(line):
            continue
        cleaned_lines.append(line)

    result = []
    blank_count = 0
    for line in cleaned_lines:
        if not line:
            blank_count += 1
            if blank_count <= 1:
                result.append("")
        else:
            blank_count = 0
            result.append(line)

    return "\n".join(result).strip()

def _extract_single_file_bytes(data: bytes, filename: str) -> str:
    extension = Path(filename).suffix.lower()
    if extension == ".pdf":
        if not data.startswith(b"%PDF-"):
            raise ResumeExtractionError("The file is not a valid PDF.")
        reader = PdfReader(BytesIO(data))
        if reader.is_encrypted:
            raise ResumeExtractionError("Password-protected PDFs are not supported.")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise ResumeExtractionError(f"PDF resumes may contain at most {MAX_PDF_PAGES} pages.")
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif extension == ".docx":
        if not data.startswith(b"PK"):
            raise ResumeExtractionError("The file is not a valid DOCX.")
        document = Document(BytesIO(data))
        blocks = [p.text for p in document.paragraphs if p.text.strip()]
        blocks.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells if cell.text.strip())
        return "\n".join(blocks)
    else:
        return data.decode("utf-8-sig")

def extract_zip_archive(data: bytes) -> list[tuple[str, str]]:
    if not data.startswith(b"PK"):
        raise ResumeExtractionError("The uploaded file is not a valid ZIP archive.")
    results = []
    try:
        with zipfile.ZipFile(BytesIO(data)) as z:
            for item in z.infolist():
                if item.is_dir() or item.filename.startswith("__MACOSX") or Path(item.filename).name.startswith("."):
                    continue
                ext = Path(item.filename).suffix.lower()
                if ext in {".pdf", ".docx", ".md", ".markdown", ".txt"}:
                    file_data = z.read(item.filename)
                    if file_data:
                        text = _extract_single_file_bytes(file_data, item.filename)
                        clean_text = _optimize_extracted_text(text)
                        if len(clean_text) >= 50:
                            results.append((Path(item.filename).stem.replace("_", " ").title(), clean_text))
                        if len(results) >= 10:
                            break
    except Exception as exc:
        raise ResumeExtractionError("Could not unpack ZIP archive.") from exc
    if not results:
        raise ResumeExtractionError("No readable PDF, DOCX, Markdown, or TXT resume files were found inside the ZIP archive.")
    return results

async def extract_resume(upload: UploadFile) -> tuple[str, str]:
    filename = Path(upload.filename or "").name
    extension = Path(filename).suffix.lower()
    if extension not in ALLOWED:
        raise ResumeExtractionError("Upload a PDF, DOCX, Markdown, TXT, or ZIP file.")
    content_type = (upload.content_type or "application/octet-stream").lower()
    if content_type not in ALLOWED[extension]:
        raise ResumeExtractionError("The file type does not match its extension.")
    data = await upload.read(MAX_FILE_BYTES + 1)
    await upload.close()
    if not data:
        raise ResumeExtractionError("The uploaded file is empty.")
    if len(data) > MAX_FILE_BYTES:
        raise ResumeExtractionError("The file must be 10 MB or smaller.")
    try:
        if extension == ".zip":
            zip_results = extract_zip_archive(data)
            text = "\n\n--- NEXT CANDIDATE ---\n\n".join(f"[{cand_name}]\n{cand_text}" for cand_name, cand_text in zip_results)
        else:
            text = _extract_single_file_bytes(data, filename)
    except ResumeExtractionError:
        raise
    except (UnicodeDecodeError, BadZipFile):
        raise ResumeExtractionError("The file could not be read. Markdown and TXT files must be UTF-8.")
    except Exception as exc:
        raise ResumeExtractionError("The file could not be read or may be damaged.") from exc
    text = _optimize_extracted_text(text)
    if len(text) < 100:
        raise ResumeExtractionError("Fewer than 100 readable characters were found. Paste the resume text instead.")
    if len(text) > MAX_TEXT_CHARS:
        raise ResumeExtractionError(f"Extracted text exceeds {MAX_TEXT_CHARS:,} characters. Shorten the resume and retry.")
    return filename, text

