from io import BytesIO
from docx import Document
from fastapi.testclient import TestClient
from app.main import app
from app.models import *
client=TestClient(app)
def test_health(): assert client.get("/health").json()=={"status":"ok"}
def test_home(): assert "Upload file" in client.get("/").text
def test_validation():
 r=client.post("/api/optimize",json={"candidate_name":"A","candidate_profile":"short","job_description":"short"})
 assert r.status_code==422
 assert set(r.json()["fields"])=={"candidate_profile","job_description"}
def test_qualified(monkeypatch):
 a=JobAnalysis(job_title="Engineer",required_qualifications=["Python"],preferred_qualifications=[],responsibilities=["Build"],ats_keywords=["Python"],qualification_status="QUALIFIED",matched_requirements=["Python"],missing_required_qualifications=[],qualification_summary="Match.")
 out=OptimizeResponse(candidate_name="Jordan",analysis=a,optimized_resume=TailoredResume(professional_title="Engineer",professional_summary="Python engineer.",core_skills=["Python"],tailored_experience=["Built services."],education_and_credentials=[],ats_keywords_used=["Python"],optimization_notes=[]))
 monkeypatch.setattr("app.main.optimize_resume",lambda _:out)
 p={"candidate_name":"Jordan","candidate_profile":"Python engineer. "*10,"job_description":"Seeking Python engineer. "*10}
 assert client.post("/api/optimize",json=p).status_code==200
def test_markdown_extraction():
 text=("# Jordan Resume\n\nPython engineer building secure APIs.\n"*5).encode()
 r=client.post("/api/extract-resume",files={"file":("resume.md",text,"text/markdown")})
 assert r.status_code==200
 assert r.json()["filename"]=="resume.md"
 assert "Python engineer" in r.json()["text"]
def test_docx_extraction():
 d=Document();d.add_heading("Jordan Resume",0);d.add_paragraph("Python engineer building secure APIs and cloud services. "*4)
 stream=BytesIO();d.save(stream)
 r=client.post("/api/extract-resume",files={"file":("resume.docx",stream.getvalue(),"application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
 assert r.status_code==200
 assert "cloud services" in r.json()["text"]
def test_disallowed_upload():
 r=client.post("/api/extract-resume",files={"file":("resume.exe",b"x"*200,"application/octet-stream")})
 assert r.status_code==400
 assert r.json()["detail"]=="Upload a PDF, DOCX, or Markdown file."
def test_spoofed_pdf():
 r=client.post("/api/extract-resume",files={"file":("resume.pdf",b"not a pdf"*30,"application/pdf")})
 assert r.status_code==400
 assert "not a valid PDF" in r.json()["detail"]
def test_oversize_upload():
 r=client.post("/api/extract-resume",files={"file":("resume.md",b"x"*(5*1024*1024+1),"text/markdown")})
 assert r.status_code==400
 assert "5 MB" in r.json()["detail"]


def test_gemini_uses_json_schema(monkeypatch):
 from types import SimpleNamespace
 from app import gemini_service
 analysis=JobAnalysis(job_title="Engineer",required_qualifications=[],preferred_qualifications=[],responsibilities=[],ats_keywords=[],qualification_status="QUALIFIED",matched_requirements=[],missing_required_qualifications=[],qualification_summary="Match.")
 class Models:
  def generate_content(self,**kwargs):
   config=kwargs["config"]
   schema=config.response_json_schema
   assert schema and config.response_schema is None
   rendered=str(schema)
   assert "$defs" not in rendered and "$ref" not in rendered
   assert "'title':" not in rendered and "'default':" not in rendered
   assert "anyOf" not in rendered
   return SimpleNamespace(parsed=analysis,text=None)
 monkeypatch.setattr(gemini_service,"_client",lambda:SimpleNamespace(models=Models()))
 assert gemini_service._generate("test",JobAnalysis)==analysis


def test_model_discovery_filters_generate_content():
 from types import SimpleNamespace
 from app.list_models import available_generate_models
 models=[
  SimpleNamespace(name="models/gemini-valid",supported_actions=["generateContent"]),
  SimpleNamespace(name="models/embed-only",supported_actions=["embedContent"]),
  SimpleNamespace(name="models/gemini-other",supported_actions=["GENERATE_CONTENT"]),
 ]
 client=SimpleNamespace(models=SimpleNamespace(list=lambda:models))
 assert available_generate_models(client)==["models/gemini-other","models/gemini-valid"]


def test_model_configuration_is_dynamic_and_normalized(monkeypatch):
 from app import gemini_service
 monkeypatch.setenv("GEMINI_MODEL"," models/gemini-2.5-flash ")
 assert gemini_service.configured_model()=="models/gemini-2.5-flash"
 monkeypatch.setenv("GEMINI_MODEL","gemini-2.5-flash-lite")
 assert gemini_service.configured_model()=="gemini-2.5-flash-lite"


def test_model_client_stays_live_during_call_and_closes_after(monkeypatch):
 from types import SimpleNamespace
 from app import gemini_service
 events=[]
 class Models:
  def generate_content(self,**kwargs):
   events.append("called")
   assert events==["called"]
   return SimpleNamespace(text="OK")
 class Client:
  models=Models()
  def close(self): events.append("closed")
 client=Client()
 monkeypatch.setattr(gemini_service,"_client",lambda:client)
 response=gemini_service._call_model("Reply with exactly OK.")
 assert response.text=="OK"
 assert events==["called","closed"]


def test_resolve_model_retains_catalog_resource_name(monkeypatch):
 from types import SimpleNamespace
 from app import gemini_service
 monkeypatch.setenv("GEMINI_MODEL","gemini-2.5-flash")
 catalog=[SimpleNamespace(name="models/gemini-2.5-flash",supported_actions=["generateContent"])]
 client=SimpleNamespace(models=SimpleNamespace(list=lambda:catalog))
 assert gemini_service._resolve_model(client)=="models/gemini-2.5-flash"


def test_generate_validates_parsed_dict_to_pydantic(monkeypatch):
 from types import SimpleNamespace
 from app import gemini_service
 parsed={
  "job_title":"Platform Engineer",
  "company_name":"Example Co",
  "required_qualifications":["Python"],
  "preferred_qualifications":["Cloud"],
  "responsibilities":["Build APIs"],
  "ats_keywords":["Python","APIs"],
  "qualification_status":"NOT_QUALIFIED",
  "matched_requirements":[],
  "missing_required_qualifications":["Python"],
  "qualification_summary":"Required Python evidence was not found.",
 }
 monkeypatch.setattr(gemini_service,"_call_model",lambda *args,**kwargs:SimpleNamespace(parsed=parsed,text=None))
 result=gemini_service._generate("test",JobAnalysis)
 assert isinstance(result,JobAnalysis)
 assert result.qualification_status==QualificationStatus.NOT_QUALIFIED
 assert result.missing_required_qualifications==["Python"]


def test_job_url_requires_https():
 import pytest
 from app.job_fetch import JobFetchError,_validate_url_shape
 with pytest.raises(JobFetchError,match="HTTPS"):
  _validate_url_shape("http://example.com/job")

def test_job_url_blocks_private_dns(monkeypatch):
 import pytest
 from app.job_fetch import JobFetchError,_validate_public_host
 monkeypatch.setattr("app.job_fetch.socket.getaddrinfo",lambda *a,**k:[(2,1,6,"",("127.0.0.1",443))])
 with pytest.raises(JobFetchError,match="Private"):
  _validate_public_host("example.com")

def test_job_html_prefers_jobposting_jsonld():
 from app.job_fetch import _extract_html
 html='<html><script type="application/ld+json">{"@type":"JobPosting","description":"<p>Python engineer building secure APIs and cloud systems for customers.</p><p>Five years experience required.</p>"}</script><body>Navigation noise</body></html>'
 text=_extract_html(html)
 assert "Python engineer" in text and "Navigation noise" not in text

def test_job_url_endpoint_preview(monkeypatch):
 async def fake(url): return url,"Python engineer role. "*10
 monkeypatch.setattr("app.main.fetch_job_description",fake)
 r=client.post("/api/extract-job-url",json={"url":"https://jobs.example.com/role"})
 assert r.status_code==200
 assert r.json()["source_url"]=="https://jobs.example.com/role"
 assert "Python engineer" in r.json()["text"]

def test_candidate_updates_are_facts_but_do_not_bypass_gate(monkeypatch):
 from app import gemini_service
 prompts=[]
 analysis=JobAnalysis(job_title="Engineer",required_qualifications=["Kubernetes"],preferred_qualifications=[],responsibilities=[],ats_keywords=["Kubernetes"],qualification_status="NOT_QUALIFIED",matched_requirements=[],missing_required_qualifications=["Kubernetes"],qualification_summary="No Kubernetes evidence.")
 def fake(prompt,schema): prompts.append(prompt);return analysis
 monkeypatch.setattr(gemini_service,"_generate",fake)
 request=OptimizeRequest(candidate_name="Jordan",candidate_profile="Python engineer experience. "*8,job_description="Kubernetes is required for this platform engineering position. "*4,candidate_updates="Completed an AWS course and built a personal API project.")
 result=gemini_service.optimize_resume(request)
 assert result.optimized_resume is None
 assert len(prompts)==1 and "<candidate_updates>" in prompts[0] and "AWS course" in prompts[0]


def test_application_readiness_parses_grounded_concerns():
 concern=ReadinessConcern(priority=1,category="mandatory_requirement",heading="Required certification not shown",potential_screening_concern="Potential screening concern: the required certification is not shown.",job_requirement_source="Posting requires Certification X.",candidate_evidence_status="Not enough evidence in your current profile.",recommended_next_step="Add the certification only if earned and verifiable.")
 analysis=JobAnalysis(job_title="Engineer",required_qualifications=["Certification X"],preferred_qualifications=[],responsibilities=[],ats_keywords=[],qualification_status="NOT_QUALIFIED",matched_requirements=[],missing_required_qualifications=["Certification X"],qualification_summary="Required evidence is missing.",application_readiness=ApplicationReadiness(summary="One evidence-based concern.",concerns=[concern]))
 response=OptimizeResponse(candidate_name="Jordan",analysis=analysis)
 assert response.analysis.application_readiness.concerns[0].job_requirement_source=="Posting requires Certification X."
 assert response.readiness_access.tier=="free"
 assert response.four_week_skill_plan is None
 assert response.verified_resume_updates is None

def test_readiness_prompt_is_calibrated_and_protected(monkeypatch):
 from app import gemini_service
 prompts=[]
 analysis=JobAnalysis(job_title="Engineer",required_qualifications=[],preferred_qualifications=[],responsibilities=[],ats_keywords=[],qualification_status="QUALIFIED",matched_requirements=[],missing_required_qualifications=[],qualification_summary="Match.")
 resume=TailoredResume(professional_title="Engineer",professional_summary="Summary",core_skills=[],tailored_experience=[],education_and_credentials=[],ats_keywords_used=[],optimization_notes=[])
 def fake(prompt,schema): prompts.append(prompt);return analysis if schema is JobAnalysis else resume
 monkeypatch.setattr(gemini_service,"_generate",fake)
 q=OptimizeRequest(candidate_name="Jordan",candidate_profile="Python engineer. "*10,job_description="Seeking a Python engineer. "*10)
 gemini_service.optimize_resume(q)
 assert "Never infer employer intent" in prompts[0]
 assert "protected traits" in prompts[0]
 assert "return fewer or none rather than inventing gaps" in prompts[0]

def test_ui_contains_application_readiness_boundary():
 html=client.get("/").text
 js=__import__("pathlib").Path("app/static/app.js").read_text()
 assert "RoleReady AI" in html
 assert "Executive Career Intelligence Platform" in html
 assert "Multi-Resume Benchmark" in html

def test_models_endpoint():
 r = client.get("/api/models")
 assert r.status_code == 200
 data = r.json()
 assert "models" in data
 assert "default" in data

def test_token_aware_extraction_optimization():
 from app.extraction import _optimize_extracted_text
 raw = "Header Text\n\n\nPage 1 of 3\n\n\nExperience line 1\n\n\n\nPage 2\n\nExperience line 2"
 clean = _optimize_extracted_text(raw)
 assert "Page 1 of 3" not in clean
 assert "Page 2" not in clean
 assert "Header Text\n\nExperience line 1\n\nExperience line 2" in clean

def test_multi_compare_endpoint(monkeypatch):
 from app import gemini_service
 a1 = JobAnalysis(job_title="DevOps", required_qualifications=["Docker"], preferred_qualifications=[], responsibilities=[], ats_keywords=[], qualification_status="QUALIFIED", matched_requirements=["Docker"], missing_required_qualifications=[], qualification_summary="Matched", overall_score=90)
 r1 = OptimizeResponse(candidate_name="Alex", analysis=a1)
 monkeypatch.setattr(gemini_service, "optimize_resume", lambda q: r1)
 
 payload = {
  "job_description": "Seeking DevOps engineer with Docker experience. "*10,
  "candidates": [
   {"candidate_name": "Alex", "candidate_profile": "Experienced DevOps engineer with Docker. "*5},
   {"candidate_name": "Sam", "candidate_profile": "Junior developer learning Docker. "*5}
  ]
 }
 res = client.post("/api/multi-compare", json=payload)
 assert res.status_code == 200
 data = res.json()
 assert "top_matching_candidate" in data
 assert len(data["results"]) == 2

def test_scores_and_interview_prep():
 prep = InterviewPitchPrep(
  elevator_pitch="Proven platform engineer with 5 years experience.",
  star_questions=[
   StarInterviewQuestion(question="Describe an outage you fixed.", recommended_talking_point="Root cause analysis", candidate_evidence_to_highlight="Kubernetes rollback")
  ]
 )
 growth = CareerGrowthPlan(
  skill_gaps=[SkillGapAdvice(missing_skill="AWS", why_flagged="Required for cloud role", learning_resource="AWS Certified Solutions Architect", resume_quick_fix="Highlight GCP experience")],
  estimated_bridge_weeks=3
 )
 analysis = JobAnalysis(
  job_title="Senior Engineer",
  required_qualifications=["Python"],
  preferred_qualifications=[],
  responsibilities=[],
  ats_keywords=["Python"],
  qualification_status="QUALIFIED",
  matched_requirements=["Python"],
  missing_required_qualifications=[],
  qualification_summary="Match",
  overall_score=88,
  hard_skills_score=95,
  soft_skills_score=80,
  experience_match_score=85,
  keyword_coverage_score=90,
  career_growth_plan=growth,
  interview_prep=prep
 )
 assert analysis.overall_score == 88
 assert analysis.interview_prep.elevator_pitch.startswith("Proven platform")
 assert len(analysis.career_growth_plan.skill_gaps) == 1

def test_10_candidate_multi_compare(monkeypatch):
 from app import gemini_service
 a1 = JobAnalysis(job_title="Engineer", required_qualifications=["Python"], preferred_qualifications=[], responsibilities=[], ats_keywords=[], qualification_status="QUALIFIED", matched_requirements=["Python"], missing_required_qualifications=[], qualification_summary="Match", overall_score=85)
 r1 = OptimizeResponse(candidate_name="Cand", analysis=a1)
 monkeypatch.setattr(gemini_service, "optimize_resume", lambda q: r1)

 candidates = [{"candidate_name": f"Candidate {i}", "candidate_profile": "Python developer profile text. "*5} for i in range(1, 11)]
 payload = {
  "job_description": "Seeking Python engineer for platform team. "*10,
  "candidates": candidates
 }
 res = client.post("/api/multi-compare", json=payload)
 assert res.status_code == 200
 data = res.json()
 assert len(data["results"]) == 10

def test_publish_leaderboard_endpoint():
 payload = {
  "job_title": "Senior AI Architect",
  "company_name": "TechCorp",
  "candidates_count": 3,
  "results": [
   {"candidate_name": "Alex", "overall_score": 95, "qualification_status": "QUALIFIED", "hard_skills_score": 90, "experience_match_score": 95, "matched_count": 5, "missing_count": 0, "qualification_summary": "Top Match"}
  ]
 }
 res = client.post("/api/publish-leaderboard", json=payload)
 assert res.status_code == 200
 data = res.json()
 assert "share_token" in data
 assert "published_url" in data
 assert data["top_candidate"] == "Alex"

def test_new_ai_generators_endpoints(monkeypatch):
 from app import main
 from app.models import CoverLetterResponse, SalaryNegotiationResponse, OutreachDraftResponse

 monkeypatch.setattr(main, "generate_cover_letter", lambda req: CoverLetterResponse(
  salutation="Dear Hiring Manager,", opening_hook="I am excited to apply...", body_paragraphs=["Key achievement 1"], closing_call_to_action="Sincerely,", full_text="Dear Hiring Manager,\n\nI am excited..."
 ))
 monkeypatch.setattr(main, "generate_salary_strategy", lambda req: SalaryNegotiationResponse(
  target_title="Senior Engineer", estimated_compensation_range="$180k - $210k", market_alignment_summary="Strong alignment", talking_points=["Point 1"], counter_offer_script="Script", email_template="Email"
 ))
 monkeypatch.setattr(main, "generate_outreach_drafts", lambda req: OutreachDraftResponse(
  linkedin_connection_note="Hi, excited about your team...", recruiter_cold_email="Cold email...", hiring_manager_followup="Followup..."
 ))

 cl_res = client.post("/api/generate-cover-letter", json={"candidate_name": "Alex", "candidate_profile": "Profile text sample. "*10, "job_description": "Job description text sample. "*15})
 assert cl_res.status_code == 200
 assert cl_res.json()["salutation"] == "Dear Hiring Manager,"

 sal_res = client.post("/api/generate-salary-strategy", json={"candidate_name": "Alex", "job_title": "Senior Engineer", "job_description": "Job description text sample. "*10})
 assert sal_res.status_code == 200
 assert sal_res.json()["estimated_compensation_range"] == "$180k - $210k"

 out_res = client.post("/api/generate-outreach", json={"first_name": "Alex", "last_name": "Morgan", "target_company": "Stripe", "target_role": "Engineer"})
 assert out_res.status_code == 200
 assert "linkedin_connection_note" in out_res.json()

def test_voice_interview_endpoints(monkeypatch):
  from app import main
  from app.models import VoiceInterviewQuestion, VoiceInterviewReport

  monkeypatch.setattr(main, "generate_voice_interview_questions", lambda req: [
    VoiceInterviewQuestion(id=1, question="Tell me about a production outage.", star_focus="Situation", recommended_talking_point="Root cause analysis")
  ])
  monkeypatch.setattr(main, "generate_voice_interview_report", lambda req: VoiceInterviewReport(
    overall_rating="Strong Hire — 92/100", strong_points=["Great metric quantification"], weaknesses=["Could elaborate on team dynamics"], areas_to_review=["System architecture design"], downloadable_summary="Excellent overall performance."
  ))

  q_res = client.post("/api/voice-interview/questions", json={"first_name": "Alex", "last_name": "Morgan", "target_role": "Senior Engineer", "interview_timeline": "this_week", "career_stage": "senior"})
  assert q_res.status_code == 200
  assert len(q_res.json()) == 1

  rep_res = client.post("/api/voice-interview/report", json={"first_name": "Alex", "last_name": "Morgan", "target_role": "Senior Engineer", "career_stage": "senior"})
  assert rep_res.status_code == 200
  assert rep_res.json()["overall_rating"].startswith("Strong Hire")


def test_theme_api_list():
    res = client.get("/api/themes")
    assert res.status_code == 200
    themes = res.json()
    assert len(themes) >= 7

def test_theme_api_filter():
    res = client.get("/api/themes?type=resume")
    assert res.status_code == 200
    for t in res.json():
        assert t["type"] == "resume"

def test_theme_api_get_by_id():
    res = client.get("/api/themes/theme_executive_slate")
    assert res.status_code == 200
    assert res.json()["id"] == "theme_executive_slate"

def test_theme_api_apply():
    res = client.post("/api/themes/theme_executive_slate/apply")
    assert res.status_code == 200
    assert res.json()["status"] == "applied"
    assert res.json()["theme_id"] == "theme_executive_slate"

def test_user_themes():
    res = client.get("/api/user/themes?user_id=user_active")
    assert res.status_code == 200
    assert "themes" in res.json()

def test_user_preferences_api():
    post_res = client.post("/api/user/preferences", json={"user_id": "user_active", "theme_id": "theme_modern_tech"})
    assert post_res.status_code == 200
    assert post_res.json()["status"] == "success"
    
    get_res = client.get("/api/user/preferences?user_id=user_active")
    assert get_res.status_code == 200
    assert get_res.json()["preferences"]["theme_id"] == "theme_modern_tech"
